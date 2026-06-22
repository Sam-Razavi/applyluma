import re
from pathlib import Path

import pdfplumber
from docx import Document

_CID_RE = re.compile(r"\(cid:\d+\)")


def _clean_extracted_text(text: str) -> str:
    return _CID_RE.sub("", text)


def parse_pdf(file_path: Path) -> str:
    with pdfplumber.open(file_path) as pdf:
        pages = [page.extract_text() or "" for page in pdf.pages]
    raw = "\n".join(pages).strip()
    return _clean_extracted_text(raw)


def parse_docx(file_path: Path) -> str:
    doc = Document(str(file_path))
    parts: list[str] = []

    for para in doc.paragraphs:
        if para.text.strip():
            parts.append(para.text.strip())

    for table in doc.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                parts.append(" | ".join(cells))

    return "\n".join(parts).strip()


def parse_cv(file_path: Path, extension: str) -> str:
    """Parse a CV file and return its text content.

    Args:
        file_path: Path to the stored file.
        extension: Normalised extension, either '.pdf' or '.docx'.

    Raises:
        ValueError: If the extension is not supported.
        Exception: Propagates parsing errors to the caller for HTTP handling.
    """
    if extension == ".pdf":
        return parse_pdf(file_path)
    if extension == ".docx":
        return parse_docx(file_path)
    raise ValueError(f"Unsupported extension: {extension}")
